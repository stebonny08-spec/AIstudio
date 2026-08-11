"""
core/markdown_converter.py
---------------------------
Converte immagini, PDF e documenti Word in file Markdown ottimizzati per AI.
Tutto elaborato in locale usando OpenCV, Tesseract OCR e librerie native.
"""

import cv2
import numpy as np
import pytesseract
from pathlib import Path
from typing import Optional, List, Tuple
import logging
import re

# Import per PDF
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("PyMuPDF non installato. Supporto PDF disabilitato.")

# Import per DOCX
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx non installato. Supporto DOCX disabilitato.")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class MarkdownConverter:
    """Classe per convertire vari formati in Markdown ottimizzato per AI"""
    
    def __init__(self, tesseract_cmd: Optional[str] = None):
        """
        Inizializza il convertitore
        
        Args:
            tesseract_cmd: Percorso personalizzato per tesseract (opzionale)
        """
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        
        # Verifica che tesseract sia installato
        try:
            pytesseract.get_tesseract_version()
            logger.info("Tesseract OCR inizializzato per MarkdownConverter")
        except Exception as e:
            logger.error(f"Tesseract non trovato: {e}")
            raise RuntimeError(
                "Tesseract OCR non è installato. "
                "Installa con: sudo apt-get install tesseract-ocr (Linux)"
            )
    
    def preprocess_image_for_ocr(self, image: np.ndarray) -> Tuple[np.ndarray, str]:
        """
        Pre-elabora un'immagine ed estrae il testo con OCR avanzato
        Ottimizzato per testo manoscritto e layout complessi
        
        Args:
            image: Immagine come array numpy
            
        Returns:
            Tuple (immagine_pre_elaborata, testo_estratto)
        """
        # Converti in scala di grigi se necessario
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Migliora contrasto con CLAHE
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        
        # Denoising avanzato
        denoised = cv2.fastNlMeansDenoising(enhanced, None, h=15, templateWindowSize=7, searchWindowSize=21)
        
        # Rileva blocchi di testo
        text_blocks = self._detect_text_blocks(denoised)
        
        if text_blocks:
            full_text = []
            
            for block in sorted(text_blocks, key=lambda b: (b['y'], b['x'])):
                x, y, w, h = block['x'], block['y'], block['w'], block['h']
                
                # Estrai ROI con margine
                margin = 5
                roi_x = max(0, x - margin)
                roi_y = max(0, y - margin)
                roi_w = min(denoised.shape[1] - roi_x, w + 2 * margin)
                roi_h = min(denoised.shape[0] - roi_y, h + 2 * margin)
                
                roi = denoised[roi_y:roi_y+roi_h, roi_x:roi_x+roi_w]
                
                # Thresholding adattivo
                roi_thresh = cv2.adaptiveThreshold(
                    roi, 255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    blockSize=11,
                    C=5
                )
                
                # OCR con configurazione ottimizzata
                try:
                    custom_config = r'--oem 3 --psm 6'
                    text = pytesseract.image_to_string(roi_thresh, lang='ita+eng', config=custom_config)
                    full_text.append(text.strip())
                except Exception as e:
                    logger.warning(f"OCR fallito su blocco: {e}")
            
            extracted_text = '\n\n'.join(full_text)
            return denoised, extracted_text
        
        else:
            # Fallback: processa l'intera immagine
            thresh = cv2.adaptiveThreshold(
                denoised, 255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                blockSize=11,
                C=5
            )
            
            try:
                custom_config = r'--oem 3 --psm 6'
                text = pytesseract.image_to_string(thresh, lang='ita+eng', config=custom_config)
            except Exception as e:
                logger.warning(f"OCR fallito: {e}")
                text = ""
            
            return thresh, text.strip()
    
    def _detect_text_blocks(self, image: np.ndarray) -> List[dict]:
        """Rileva blocchi di testo nell'immagine"""
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Soglia adattiva
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Trova contorni
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        blocks = []
        min_area = 500
        
        for contour in contours:
            area = cv2.contourArea(contour)
            if area > min_area:
                x, y, w, h = cv2.boundingRect(contour)
                if w > 20 and h > 20:
                    blocks.append({'x': x, 'y': y, 'w': w, 'h': h, 'area': area})
        
        return blocks
    
    def _format_as_markdown(self, text: str, source_type: str = "generic") -> str:
        """
        Formatta il testo estratto in Markdown strutturato e leggibile da AI
        
        Args:
            text: Testo estratto
            source_type: Tipo di sorgente ("image", "pdf", "docx")
            
        Returns:
            Stringa formattata in Markdown
        """
        if not text:
            return ""
        
        # Dividi in paragrafi
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        markdown_lines = []
        markdown_lines.append(f"# Appunti Convertiti ({source_type.upper()})\n")
        markdown_lines.append(f"*Convertito automaticamente da StudioIA*\n")
        markdown_lines.append("---\n")
        
        current_section = None
        in_list = False
        
        for para in paragraphs:
            # Rileva titoli (parole corte seguite da due punti o tutto maiuscolo)
            if len(para) < 80 and (para.endswith(':') or para.isupper()):
                current_section = para.rstrip(':')
                markdown_lines.append(f"\n## {current_section}\n")
                in_list = False
            # Rileva liste (linee che iniziano con -, *, •, numeri)
            elif re.match(r'^[\-\*•]\s+', para) or re.match(r'^\d+[\.\\)\s]', para):
                if not in_list:
                    markdown_lines.append("")
                    in_list = True
                markdown_lines.append(f"- {para.lstrip('-*•').strip()}")
            # Paragrafo normale
            else:
                if in_list:
                    markdown_lines.append("")
                    in_list = False
                markdown_lines.append(para)
        
        markdown_lines.append("\n---\n*Fine del documento*")
        
        return '\n'.join(markdown_lines)
    
    def convert_image_to_markdown(self, image_path: str, output_path: Optional[str] = None) -> str:
        """
        Converte un'immagine in file Markdown
        
        Args:
            image_path: Percorso dell'immagine
            output_path: Percorso opzionale per il file .md di output
            
        Returns:
            Contenuto Markdown generato
        """
        image = cv2.imread(str(image_path))
        if image is None:
            raise ValueError(f"Impossibile caricare l'immagine: {image_path}")
        
        _, text = self.preprocess_image_for_ocr(image)
        markdown_content = self._format_as_markdown(text, "image")
        
        if output_path:
            output_file = Path(output_path)
            if not output_file.suffix == '.md':
                output_file = output_file.with_suffix('.md')
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            logger.info(f"Markdown salvato in: {output_file}")
        
        return markdown_content
    
    def convert_pdf_to_markdown(self, pdf_path: str, output_path: Optional[str] = None) -> str:
        """
        Converte un PDF in file Markdown
        
        Args:
            pdf_path: Percorso del PDF
            output_path: Percorso opzionale per il file .md di output
            
        Returns:
            Contenuto Markdown generato
        """
        if not PDF_SUPPORT:
            raise RuntimeError("PyMuPDF non installato. Installa con: pip install PyMuPDF")
        
        pdf_file = Path(pdf_path)
        if not pdf_file.exists():
            raise FileNotFoundError(f"PDF non trovato: {pdf_path}")
        
        all_text = []
        
        try:
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Estrai testo nativo se disponibile
                text = page.get_text()
                
                # Se poco testo, usa OCR sulla pagina renderizzata
                if len(text.strip()) < 50:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    img_array = np.frombuffer(img_data, np.uint8).reshape((pix.height, pix.width, 4))
                    img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGBA2BGR)
                    _, ocr_text = self.preprocess_image_for_ocr(img_bgr)
                    text = ocr_text
                
                if text.strip():
                    all_text.append(f"### Pagina {page_num + 1}\n\n{text}")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"Errore lettura PDF: {e}")
            raise
        
        full_text = '\n\n'.join(all_text)
        markdown_content = self._format_as_markdown(full_text, "pdf")
        
        if output_path:
            output_file = Path(output_path)
            if not output_file.suffix == '.md':
                output_file = output_file.with_suffix('.md')
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            logger.info(f"Markdown salvato in: {output_file}")
        
        return markdown_content
    
    def convert_docx_to_markdown(self, docx_path: str, output_path: Optional[str] = None) -> str:
        """
        Converte un documento Word in file Markdown
        
        Args:
            docx_path: Percorso del file .docx
            output_path: Percorso opzionale per il file .md di output
            
        Returns:
            Contenuto Markdown generato
        """
        if not DOCX_SUPPORT:
            raise RuntimeError("python-docx non installato. Installa con: pip install python-docx")
        
        docx_file = Path(docx_path)
        if not docx_file.exists():
            raise FileNotFoundError(f"Documento non trovato: {docx_path}")
        
        try:
            doc = Document(docx_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Rileva stile del paragrafo
                    style_name = para.style.name.lower() if para.style else ""
                    
                    if 'heading 1' in style_name:
                        paragraphs.append(f"# {text}")
                    elif 'heading 2' in style_name:
                        paragraphs.append(f"## {text}")
                    elif 'heading 3' in style_name:
                        paragraphs.append(f"### {text}")
                    elif 'list' in style_name or text.startswith(('-', '*', '•')):
                        paragraphs.append(f"- {text.lstrip('-*•').strip()}")
                    else:
                        paragraphs.append(text)
            
            # Aggiungi tabelle se presenti
            for table in doc.tables:
                paragraphs.append("\n| " + " | ".join([cell.text for cell in table.rows[0].cells]) + " |")
                paragraphs.append("|" + "|".join(["---" for _ in table.rows[0].cells]) + "|")
                
                for row in table.rows[1:]:
                    paragraphs.append("| " + " | ".join([cell.text for cell in row.cells]) + " |")
        
        except Exception as e:
            logger.error(f"Errore lettura DOCX: {e}")
            raise
        
        full_text = '\n\n'.join(paragraphs)
        markdown_content = self._format_as_markdown(full_text, "docx")
        
        if output_path:
            output_file = Path(output_path)
            if not output_file.suffix == '.md':
                output_file = output_file.with_suffix('.md')
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            logger.info(f"Markdown salvato in: {output_file}")
        
        return markdown_content
    
    def convert_file_to_markdown(self, file_path: str, output_folder: str) -> str:
        """
        Converte automaticamente un file in Markdown in base al tipo
        
        Args:
            file_path: Percorso del file da convertire
            output_folder: Cartella dove salvare il file .md
            
        Returns:
            Percorso del file Markdown creato
        """
        file_path = Path(file_path)
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)
        
        output_filename = file_path.stem + ".md"
        output_path = output_folder / output_filename
        
        suffix = file_path.suffix.lower()
        
        if suffix in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            self.convert_image_to_markdown(str(file_path), str(output_path))
        elif suffix == '.pdf':
            self.convert_pdf_to_markdown(str(file_path), str(output_path))
        elif suffix == '.docx':
            self.convert_docx_to_markdown(str(file_path), str(output_path))
        else:
            raise ValueError(f"Formato non supportato: {suffix}")
        
        return str(output_path)


# Funzione utility per uso rapido
def convert_to_markdown(input_path: str, output_folder: str) -> str:
    """
    Funzione rapida per convertire un file in Markdown
    
    Args:
        input_path: Percorso del file da convertire
        output_folder: Cartella dove salvare il file .md
        
    Returns:
        Percorso del file Markdown creato
    """
    converter = MarkdownConverter()
    return converter.convert_file_to_markdown(input_path, output_folder)
