"""
core/parsers/docx_parser.py
Estrazione da file Word (.docx): testo dei paragrafi, testo delle tabelle,
e immagini incorporate collegate al testo dei paragrafi circostanti.
"""

from docx import Document

from core.models import ExtractedImage, ParsedDocument

# Namespace XML usati dentro un file .docx (che è internamente uno ZIP/XML)
_NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _find_blip_rids(run_element) -> list:
    """Cerca dentro un run di Word eventuali immagini incorporate e ritorna
    gli id di relazione (rId) usati per recuperare i byte reali dell'immagine.
    """
    rids = []
    for blip in run_element.findall(".//" + _NS_A + "blip"):
        rid = blip.get(_NS_R + "embed")
        if rid:
            rids.append(rid)
    return rids


def extract(file_path: str) -> ParsedDocument:
    try:
        document = Document(file_path)
    except Exception as e:
        return ParsedDocument(source_file=file_path, error=f"Word illeggibile: {e}")

    paragraphs_text = [p.text for p in document.paragraphs]
    full_text = "\n".join(t for t in paragraphs_text if t.strip())

    # Testo delle tabelle (spesso contengono informazioni importanti)
    table_text_parts = []
    try:
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_text_parts.append(" | ".join(cells))
    except Exception:
        pass
    if table_text_parts:
        full_text += "\n\n[Tabelle]\n" + "\n".join(table_text_parts)

    # Immagini incorporate, collegate ai paragrafi vicini per capire il contesto
    images = []
    for idx, paragraph in enumerate(document.paragraphs):
        rids = []
        try:
            for run in paragraph.runs:
                rids.extend(_find_blip_rids(run._element))
        except Exception:
            continue
        if not rids:
            continue

        window = paragraphs_text[max(0, idx - 2): idx + 3]
        nearby = "\n".join(t for t in window if t.strip())[:1500] or full_text[:1500]

        for img_idx, rid in enumerate(rids):
            try:
                part = document.part.rels[rid].target_part
                images.append(
                    ExtractedImage(
                        source_file=file_path,
                        location_hint=f"paragrafo {idx + 1}, immagine {img_idx + 1}",
                        image_bytes=part.blob,
                        mime_type=part.content_type,
                        nearby_text=nearby,
                    )
                )
            except (KeyError, AttributeError):
                # Relazione non trovata o parte non valida: saltiamo la
                # singola immagine, il resto del documento resta comunque utile.
                continue

    return ParsedDocument(source_file=file_path, text=full_text, images=images)
